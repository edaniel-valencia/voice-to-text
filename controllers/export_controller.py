import io
from datetime import datetime
from docx import Document
from fpdf import FPDF

def export_txt(text: str) -> bytes:
    """Export text to a plain TXT format byte string."""
    return text.encode("utf-8")

def export_all_txt(history_rows: list) -> bytes:
    """Combines all transcriptions into a single TXT format byte string."""
    lines = ["=== HISTORIAL COMPLETO DE TRANSCRIPCIONES ===\n"]
    for row in history_rows:
        row_id, filename, language, model_used, text, word_count, created_at = row
        lines.append(f"[{created_at}] Archivo: {filename} | Idioma: {language} | Modelo: {model_used}")
        lines.append(text)
        lines.append("-" * 60 + "\n")
    return "\n".join(lines).encode("utf-8")

def export_docx(text: str, filename: str = "Transcripción") -> bytes:
    """Export text to a DOCX format byte string."""
    doc = Document()
    doc.add_heading("VoiceScript — Transcripción", level=0)
    doc.add_paragraph(f"Archivo: {filename}")
    doc.add_paragraph(f"Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph("")
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

from pathlib import Path

def export_pdf(text: str, filename: str = "Transcripción") -> bytes:
    """Export text to a PDF format byte string."""
    pdf = FPDF()
    pdf.set_margins(20, 20, 20)
    pdf.add_page()
    
    font_path = Path(__file__).parent.parent / "NotoSans-Regular.ttf"
    if font_path.exists():
        pdf.add_font("NotoSans", "", str(font_path), uni=True)
        pdf.set_font("NotoSans", "", 20)
    else:
        pdf.set_font("Helvetica", "B", 20)
        
    pdf.set_text_color(124, 58, 237)
    pdf.cell(0, 12, "VoiceScript", ln=True)
    
    if font_path.exists():
        pdf.set_font("NotoSans", "", 10)
    else:
        pdf.set_font("Helvetica", "", 10)
        
    pdf.set_text_color(100, 116, 139)
    pdf.cell(0, 6, f"Archivo: {filename}   |   Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M')}", ln=True)
    pdf.ln(6)
    pdf.set_draw_color(124, 58, 237)
    pdf.set_line_width(0.5)
    pdf.line(20, pdf.get_y(), 190, pdf.get_y())
    pdf.ln(8)
    
    if font_path.exists():
        pdf.set_font("NotoSans", "", 12)
    else:
        pdf.set_font("Helvetica", "", 12)
        
    pdf.set_text_color(30, 30, 50)
    pdf.set_auto_page_break(auto=True, margin=20)
    
    try:
        pdf.multi_cell(0, 7, text)
    except Exception as e:
        pdf.multi_cell(0, 7, f"[Error de codificación en PDF: El texto contiene caracteres no soportados].\n\nDetalle: {e}")
        
    return bytes(pdf.output())
