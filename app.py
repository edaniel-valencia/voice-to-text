import streamlit as st
import whisper
import tempfile
import os
import io
import time
import sqlite3
from typing import Optional
from pathlib import Path
from datetime import datetime
from pydub import AudioSegment
from docx import Document
from fpdf import FPDF
from openai import OpenAI
from dotenv import load_dotenv

# ─── Cargar variables de entorno ──────────────────────────────────────────────
load_dotenv()

OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434/v1")
OLLAMA_API_KEY  = os.getenv("OLLAMA_API_KEY",  "ollama")
OLLAMA_MODEL    = os.getenv("OLLAMA_MODEL",    "llama3.2")

@st.cache_resource(show_spinner=False)
def get_ollama_client():
    return OpenAI(base_url=OLLAMA_BASE_URL, api_key=OLLAMA_API_KEY)

OLLAMA_PROMPTS = {
    "✨ Mejorar redacción":   "Mejora la redacción del siguiente texto transcrito de audio. Corrige la gramática, añade puntuación correcta y haz el texto más fluido. Devuelve SOLO el texto mejorado, sin explicaciones:\n\n",
    "📋 Resumir":             "Resume el siguiente texto en 3-5 puntos clave concisos. Devuelve SOLO el resumen:\n\n",
    "🔤 Corregir puntuación": "Añade puntuación y mayúsculas correctas al siguiente texto. Devuelve SOLO el texto corregido:\n\n",
    "📝 Formato lista":       "Convierte el siguiente texto en una lista de puntos ordenados y claros. Devuelve SOLO la lista:\n\n",
    "🌐 Traducir al inglés":  "Translate the following text to English. Return ONLY the translation:\n\n",
}

def run_ollama(prompt_prefix: str, text: str) -> str:
    """Envía texto a Ollama y devuelve la respuesta."""
    try:
        client = get_ollama_client()
        response = client.chat.completions.create(
            model=OLLAMA_MODEL,
            messages=[{"role": "user", "content": prompt_prefix + text}],
            temperature=0.3,
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"❌ Error con Ollama: {e}"

# ─── Page Config ──────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="ADAVAM — VoiceScript",
    page_icon="🎙️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ─── Premium CSS ───────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800;900&display=swap');

html, body, [class*="css"] {
    font-family: 'Inter', sans-serif !important;
}

/* Background */
.stApp {
    background: linear-gradient(135deg, #0a0a18 0%, #0f0f2d 40%, #12122a 100%);
    min-height: 100vh;
}

/* Ocultar solo lo innecesario de Streamlit */
#MainMenu { visibility: hidden; }
footer { visibility: hidden; }
.stDeployButton { display: none !important; }

/* ── Sidebar SIEMPRE abierto ── */
/* Ocultar TODOS los botones de colapsar */
[data-testid="collapsedControl"],
[data-testid="stSidebarCollapseButton"],
button[aria-label="Close sidebar"],
button[aria-label="Collapse sidebar"],
button[aria-label="open sidebar"],
button[title="Collapse sidebar"] { display: none !important; }


/* ── Sidebar estilo ── */
[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #0d0d24 0%, #0f0f2d 100%) !important;
    border-right: 1px solid rgba(124, 58, 237, 0.25) !important;
}
[data-testid="stSidebar"] * { color: #cbd5e1 !important; }

/* ── Ajuste equilibrado de espacios en el sidebar ── */
[data-testid="stSidebar"] [data-testid="stVerticalBlock"] {
    gap: 16px !important;
}
/* Separadores equilibrados */
[data-testid="stSidebar"] hr {
    margin: 12px 0 !important;
    border-color: rgba(124, 58, 237, 0.15) !important;
}
/* Métricas más cómodas */
[data-testid="stSidebar"] [data-testid="stMetric"] {
    padding: 12px 14px !important;
    margin-bottom: 8px !important;
}
[data-testid="stSidebar"] .stSelectbox {
    margin-bottom: 4px !important;
}

/* ── Contenido principal ocupa pantalla completa ── */
.main { min-height: 100vh !important; }
.main .block-container {
    min-height: calc(100vh - 60px) !important;
    padding-top: 1.5rem !important;
    padding-bottom: 40px !important;
    display: flex !important;
    flex-direction: column !important;
}
section.main > div { min-height: 100vh !important; }


/* ── Main title ── */
.hero-title {
    background: linear-gradient(135deg, #a78bfa 0%, #60a5fa 50%, #34d399 100%);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    font-size: 3.2rem;
    font-weight: 900;
    line-height: 1.1;
    text-align: center;
    margin: 0;
}
.hero-subtitle {
    color: #64748b;
    font-size: 1.05rem;
    text-align: center;
    margin-top: 8px;
    margin-bottom: 32px;
}

/* ── Cards ── */
.glass-card {
    background: rgba(255, 255, 255, 0.04);
    backdrop-filter: blur(20px);
    border: 1px solid rgba(255, 255, 255, 0.08);
    border-radius: 20px;
    padding: 28px;
    margin: 12px 0;
    transition: border-color 0.3s ease;
}
.glass-card:hover {
    border-color: rgba(124, 58, 237, 0.3);
}

/* ── History item ── */
.history-item {
    background: rgba(255, 255, 255, 0.03);
    border: 1px solid rgba(255, 255, 255, 0.07);
    border-radius: 14px;
    padding: 16px 20px;
    margin: 8px 0;
    cursor: pointer;
    transition: all 0.25s ease;
}
.history-item:hover {
    background: rgba(124, 58, 237, 0.1);
    border-color: rgba(124, 58, 237, 0.4);
    transform: translateX(4px);
}
.history-meta {
    color: #64748b;
    font-size: 0.78rem;
    margin-top: 4px;
}
.history-preview {
    color: #94a3b8;
    font-size: 0.9rem;
    white-space: nowrap;
    overflow: hidden;
    text-overflow: ellipsis;
}

/* ── Buttons ── */
.stButton > button {
    background: linear-gradient(135deg, #7c3aed, #4f46e5) !important;
    color: white !important;
    border: none !important;
    border-radius: 12px !important;
    padding: 10px 24px !important;
    font-weight: 600 !important;
    font-size: 14px !important;
    transition: all 0.3s ease !important;
    box-shadow: 0 4px 20px rgba(124, 58, 237, 0.35) !important;
    letter-spacing: 0.01em !important;
}
.stButton > button:hover {
    transform: translateY(-2px) !important;
    box-shadow: 0 8px 30px rgba(124, 58, 237, 0.55) !important;
}
.stButton > button:active {
    transform: translateY(0px) !important;
}

/* Download buttons */
.stDownloadButton > button {
    background: rgba(255, 255, 255, 0.07) !important;
    color: #e2e8f0 !important;
    border: 1px solid rgba(255, 255, 255, 0.12) !important;
    border-radius: 12px !important;
    font-weight: 500 !important;
    transition: all 0.3s ease !important;
    box-shadow: none !important;
}
.stDownloadButton > button:hover {
    background: rgba(124, 58, 237, 0.2) !important;
    border-color: rgba(124, 58, 237, 0.5) !important;
    transform: translateY(-2px) !important;
    box-shadow: 0 4px 16px rgba(124, 58, 237, 0.25) !important;
}

/* ── File uploader ── */
[data-testid="stFileUploader"] {
    background: rgba(124, 58, 237, 0.04) !important;
    border: 2px dashed rgba(124, 58, 237, 0.35) !important;
    border-radius: 18px !important;
    padding: 12px !important;
    transition: all 0.3s ease !important;
}
[data-testid="stFileUploader"]:hover {
    border-color: rgba(124, 58, 237, 0.7) !important;
    background: rgba(124, 58, 237, 0.08) !important;
}

/* ── iframes (Custom Components like audio recorder) ── */
iframe {
    background-color: transparent !important;
}
[data-testid="stIFrame"] {
    background-color: transparent !important;
}


/* ── Text area ── */
.stTextArea textarea {
    background: rgba(255, 255, 255, 0.04) !important;
    border: 1px solid rgba(255, 255, 255, 0.1) !important;
    border-radius: 14px !important;
    color: #e2e8f0 !important;
    font-size: 15px !important;
    line-height: 1.75 !important;
    font-family: 'Inter', sans-serif !important;
    caret-color: #a78bfa !important;
}
.stTextArea textarea:focus {
    border-color: rgba(124, 58, 237, 0.6) !important;
    box-shadow: 0 0 0 3px rgba(124, 58, 237, 0.12) !important;
}

/* ── Selectbox ── */
[data-testid="stSelectbox"] > div > div {
    background: rgba(255, 255, 255, 0.05) !important;
    border: 1px solid rgba(255, 255, 255, 0.1) !important;
    border-radius: 10px !important;
    color: #e2e8f0 !important;
}

/* ── Progress bar ── */
.stProgress > div > div {
    background: linear-gradient(90deg, #7c3aed, #60a5fa, #34d399) !important;
    border-radius: 999px !important;
    background-size: 200% 100% !important;
    animation: shimmer 1.5s infinite linear !important;
}
@keyframes shimmer {
    0% { background-position: 200% 0; }
    100% { background-position: -200% 0; }
}

/* ── Metric ── */
[data-testid="stMetric"] {
    background: rgba(255, 255, 255, 0.04) !important;
    border: 1px solid rgba(255, 255, 255, 0.08) !important;
    border-radius: 14px !important;
    padding: 16px 20px !important;
}
[data-testid="stMetricValue"] { color: #a78bfa !important; font-weight: 700 !important; }
[data-testid="stMetricLabel"] { color: #64748b !important; }

/* ── Tabs ── */
.stTabs [data-baseweb="tab-list"] {
    background: rgba(255, 255, 255, 0.03) !important;
    border-radius: 18px !important;
    padding: 6px !important;
    gap: 8px !important;
    border: 1px solid rgba(255, 255, 255, 0.08) !important;
}
.stTabs [data-baseweb="tab"] {
    border-radius: 13px !important;
    color: #64748b !important;
    font-weight: 600 !important;
    font-size: 15px !important;
    padding: 14px 32px !important;
    letter-spacing: 0.02em !important;
    transition: all 0.25s ease !important;
    min-width: 160px !important;
    text-align: center !important;
    justify-content: center !important;
}
.stTabs [data-baseweb="tab"]:hover {
    color: #a78bfa !important;
    background: rgba(124, 58, 237, 0.08) !important;
}
.stTabs [aria-selected="true"] {
    background: linear-gradient(135deg, #7c3aed, #4f46e5) !important;
    color: white !important;
    box-shadow: 0 4px 20px rgba(124, 58, 237, 0.5) !important;
}

/* ── Alerts ── */
.stSuccess {
    background: rgba(52, 211, 153, 0.08) !important;
    border: 1px solid rgba(52, 211, 153, 0.3) !important;
    border-radius: 12px !important;
    color: #34d399 !important;
}
.stInfo {
    background: rgba(96, 165, 250, 0.08) !important;
    border: 1px solid rgba(96, 165, 250, 0.3) !important;
    border-radius: 12px !important;
    color: #60a5fa !important;
}
.stWarning {
    background: rgba(251, 191, 36, 0.08) !important;
    border: 1px solid rgba(251, 191, 36, 0.3) !important;
    border-radius: 12px !important;
}
.stError {
    background: rgba(248, 113, 113, 0.08) !important;
    border: 1px solid rgba(248, 113, 113, 0.3) !important;
    border-radius: 12px !important;
}

/* ── Divider ── */
hr { border-color: rgba(255, 255, 255, 0.07) !important; }

/* ── Spinner ── */
[data-testid="stSpinner"] { color: #a78bfa !important; }

/* ── Labels ── */
.stTextArea label, .stFileUploader label, .stSelectbox label {
    color: #94a3b8 !important;
    font-weight: 500 !important;
    font-size: 0.88rem !important;
    text-transform: uppercase !important;
    letter-spacing: 0.08em !important;
}

/* ── Badge ── */
.lang-badge {
    display: inline-block;
    background: rgba(124, 58, 237, 0.15);
    border: 1px solid rgba(124, 58, 237, 0.35);
    color: #a78bfa;
    border-radius: 999px;
    padding: 3px 12px;
    font-size: 0.78rem;
    font-weight: 600;
    letter-spacing: 0.05em;
}

/* ── Section headers ── */
.section-header {
    color: #f1f5f9;
    font-size: 1.1rem;
    font-weight: 700;
    margin: 24px 0 12px;
    display: flex;
    align-items: center;
    gap: 8px;
}

/* ── Sidebar section header ── */
.sidebar-header {
    color: #a78bfa !important;
    font-size: 0.75rem !important;
    font-weight: 700 !important;
    text-transform: uppercase !important;
    letter-spacing: 0.12em !important;
    margin: 4px 0 2px !important;
}
</style>
""", unsafe_allow_html=True)


# ─── SQLite Database ──────────────────────────────────────────────────────────
DB_PATH = Path(__file__).parent / "transcriptions.db"

def init_db():
    """Create database and table if they don't exist."""
    conn = sqlite3.connect(DB_PATH)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS transcriptions (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            filename    TEXT,
            language    TEXT,
            model       TEXT,
            text        TEXT,
            word_count  INTEGER,
            created_at  TEXT
        )
    """)
    conn.commit()
    conn.close()

def save_transcription(filename, language, model, text):
    """Persist a transcription to SQLite."""
    conn = sqlite3.connect(DB_PATH)
    word_count = len(text.split())
    created_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    conn.execute(
        "INSERT INTO transcriptions (filename, language, model, text, word_count, created_at) VALUES (?, ?, ?, ?, ?, ?)",
        (filename, language, model, text, word_count, created_at)
    )
    conn.commit()
    conn.close()

def load_history(limit=30):
    """Load recent transcriptions from SQLite."""
    conn = sqlite3.connect(DB_PATH)
    rows = conn.execute(
        "SELECT id, filename, language, model, text, word_count, created_at FROM transcriptions ORDER BY id DESC LIMIT ?",
        (limit,)
    ).fetchall()
    conn.close()
    return rows

def delete_transcription(row_id):
    """Delete a transcription by id."""
    conn = sqlite3.connect(DB_PATH)
    conn.execute("DELETE FROM transcriptions WHERE id = ?", (row_id,))
    conn.commit()
    conn.close()

def clear_all_history():
    """Delete all transcriptions."""
    conn = sqlite3.connect(DB_PATH)
    conn.execute("DELETE FROM transcriptions")
    conn.commit()
    conn.close()

init_db()


# ─── Whisper Model (cached) ───────────────────────────────────────────────────
@st.cache_resource(show_spinner=False)
def load_whisper_model(model_name: str):
    return whisper.load_model(model_name)


# ─── Audio Processing ─────────────────────────────────────────────────────────
SUPPORTED_FORMATS = ["wav", "mp3", "aac", "ogg", "flac", "m4a", "mp4", "webm"]

def to_wav(input_path: str) -> str:
    """Convert any supported audio format to WAV."""
    ext = Path(input_path).suffix.lower().lstrip(".")
    if ext == "wav":
        return input_path
    audio = AudioSegment.from_file(input_path)
    wav_path = input_path.rsplit(".", 1)[0] + "_converted.wav"
    audio.export(wav_path, format="wav")
    return wav_path

def transcribe_audio(audio_path: str, model, lang_code: Optional[str]):
    """Run Whisper transcription with progress feedback."""
    progress = st.progress(0, text="")
    status = st.empty()

    status.markdown("🔄 **Procesando audio...**")
    progress.progress(15, text="Procesando audio...")
    wav_path = to_wav(audio_path)

    status.markdown("🧠 **Transcribiendo con Whisper...**")
    progress.progress(40, text="Transcribiendo con IA...")

    options = {}
    if lang_code:
        options["language"] = lang_code

    result = model.transcribe(wav_path, **options)

    progress.progress(95, text="Finalizando...")
    time.sleep(0.3)
    progress.progress(100, text="¡Completado!")
    time.sleep(0.4)

    progress.empty()
    status.empty()

    # Cleanup converted file
    if wav_path != audio_path and os.path.exists(wav_path):
        os.remove(wav_path)

    return result["text"].strip(), result.get("language", "desconocido")


# ─── Export Helpers ───────────────────────────────────────────────────────────
def export_txt(text: str) -> bytes:
    return text.encode("utf-8")

def export_docx(text: str, filename: str = "Transcripción") -> bytes:
    doc = Document()
    doc.add_heading("VoiceScript — Transcripción", level=0)
    doc.add_paragraph(f"Archivo: {filename}")
    doc.add_paragraph(f"Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph("")
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

def export_pdf(text: str, filename: str = "Transcripción") -> bytes:
    pdf = FPDF()
    pdf.set_margins(20, 20, 20)
    pdf.add_page()
    
    # Intentar cargar fuente Unicode si existe
    font_path = Path(__file__).parent / "NotoSans-Regular.ttf"
    if font_path.exists():
        pdf.add_font("NotoSans", "", str(font_path), uni=True)
        pdf.set_font("NotoSans", "", 20)
    else:
        pdf.set_font("Helvetica", "B", 20)
        
    pdf.set_text_color(124, 58, 237)
    pdf.cell(0, 12, "VoiceScript", ln=True)
    
    if font_path.exists():
        pdf.set_font("NotoSans", "", 10)
    else:
        pdf.set_font("Helvetica", "", 10)
        
    pdf.set_text_color(100, 116, 139)
    pdf.cell(0, 6, f"Archivo: {filename}   |   Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M')}", ln=True)
    pdf.ln(6)
    pdf.set_draw_color(124, 58, 237)
    pdf.set_line_width(0.5)
    pdf.line(20, pdf.get_y(), 190, pdf.get_y())
    pdf.ln(8)
    
    if font_path.exists():
        pdf.set_font("NotoSans", "", 12)
    else:
        pdf.set_font("Helvetica", "", 12)
        
    pdf.set_text_color(30, 30, 50)
    pdf.set_auto_page_break(auto=True, margin=20)
    try:
        pdf.multi_cell(0, 7, text)
    except Exception as e:
        pdf.multi_cell(0, 7, f"[Error de codificación en PDF: El texto contiene caracteres no soportados por la fuente. Usa la descarga en TXT o DOCX].\n\nDetalle: {e}")
        
    return bytes(pdf.output())


# ─── Language Map ─────────────────────────────────────────────────────────────
LANG_OPTIONS = {
    "🌐 Auto-detectar": None,
    "🇪🇸 Español":      "es",
    "🇺🇸 English":       "en",
    "🇫🇷 Français":      "fr",
    "🇧🇷 Português":     "pt",
    "🇩🇪 Deutsch":       "de",
    "🇮🇹 Italiano":      "it",
    "🇯🇵 日本語":         "ja",
    "🇨🇳 中文":           "zh",
}

MODEL_INFO = {
    "tiny":   ("~40 MB",  "Más rápido"),
    "base":   ("~150 MB", "Recomendado ✨"),
    "small":  ("~500 MB", "Mejor precisión"),
    "medium": ("~1.5 GB", "Alta precisión"),
}


# ─── Sidebar ──────────────────────────────────────────────────────────────────
with st.sidebar:
    # ── TOP: Logo ADAVAM ──
    st.image("adavam_logo.png", use_container_width=True)
    st.markdown(
        '<p style="color:#64748b; font-size:0.75rem; text-align:center; '
        'margin:-4px 0 4px; letter-spacing:0.12em;">VOICE TO TEXT · IA</p>',
        unsafe_allow_html=True
    )
    st.markdown("---")

    # ── MIDDLE: Configuración ──
    st.markdown('<p class="sidebar-header">⚙️ Modelo de IA</p>', unsafe_allow_html=True)
    model_choice = st.selectbox(
        "Modelo",
        list(MODEL_INFO.keys()),
        index=1,
        format_func=lambda m: f"{m}  —  {MODEL_INFO[m][0]}  ({MODEL_INFO[m][1]})",
        label_visibility="collapsed"
    )
    st.caption(f"💡 **{MODEL_INFO[model_choice][0]}** · {MODEL_INFO[model_choice][1]}")
    st.markdown("---")

    st.markdown('<p class="sidebar-header">🌐 Idioma de transcripción</p>', unsafe_allow_html=True)
    lang_choice = st.selectbox(
        "Idioma",
        list(LANG_OPTIONS.keys()),
        label_visibility="collapsed"
    )
    selected_lang = LANG_OPTIONS[lang_choice]
    st.markdown("---")

    st.markdown('<p class="sidebar-header">📊 Estadísticas</p>', unsafe_allow_html=True)
    history = load_history()
    total_words = sum(r[5] for r in history) if history else 0
    col_a, col_b = st.columns(2)
    col_a.metric("📝 Total", len(history))
    col_b.metric("🔤 Palabras", f"{total_words:,}")

    st.markdown("---")
    if st.button("🗑️ Borrar todo el historial", use_container_width=True):
        clear_all_history()
        st.rerun()
    st.caption("Base de datos: `transcriptions.db`")
    st.caption(f"© {datetime.now().year} ADAVAM · VoiceScript")







# ─── Hero Header ─────────────────────────────────────────────────────────────
st.markdown("""
<h1 class="hero-title">🎙️ VoiceScript</h1>
<p class="hero-subtitle">
    Transcriptor de voz a texto con IA &nbsp;·&nbsp;
    Potenciado por Whisper &nbsp;·&nbsp;
    Historial persistente
</p>
""", unsafe_allow_html=True)



# ─── Main Tabs ────────────────────────────────────────────────────────────────
tab_upload, tab_record, tab_history = st.tabs(
    ["📁  Subir Archivo", "🎤  Grabar Audio", "🕓  Historial"]
)


# ══════════════════════════════════════════════════════════════════════════════
# TAB 1 — Upload File
# ══════════════════════════════════════════════════════════════════════════════
with tab_upload:
    st.markdown("")
    uploaded = st.file_uploader(
        "Arrastra un archivo de audio aquí o haz clic para seleccionar",
        type=SUPPORTED_FORMATS,
        help="Formatos soportados: WAV · MP3 · AAC · OGG · FLAC · M4A"
    )

    if uploaded:
        st.audio(uploaded)
        st.markdown("")

        col_l, col_btn, col_r = st.columns([1, 2, 1])
        with col_btn:
            run_btn = st.button("🚀 Transcribir", use_container_width=True, key="btn_upload")

        if run_btn:
            with st.spinner("Cargando modelo de IA…"):
                model = load_whisper_model(model_choice)

            suffix = Path(uploaded.name).suffix or ".wav"
            with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
                tmp.write(uploaded.read())
                tmp_path = tmp.name

            try:
                text, detected = transcribe_audio(tmp_path, model, selected_lang)
                save_transcription(uploaded.name, detected, model_choice, text)
                st.session_state["last_text"]     = text
                st.session_state["last_filename"] = uploaded.name
                st.session_state["last_lang"]     = detected
                st.success(f"✅ Transcripción guardada · Idioma detectado: **{detected}**")
            except Exception as e:
                st.error(f"❌ Error durante la transcripción: {e}")
            finally:
                if os.path.exists(tmp_path):
                    os.remove(tmp_path)


# ══════════════════════════════════════════════════════════════════════════════
# TAB 2 — Record Audio
# ══════════════════════════════════════════════════════════════════════════════
with tab_record:
    st.markdown("")
    try:
        from audio_recorder_streamlit import audio_recorder

        st.info("🎙️ **Haz clic en el botón de micrófono para comenzar a grabar.** Haz clic de nuevo para detener. El audio se procesará automáticamente.")

        st.markdown("""
        <div style="text-align:center; margin: 28px 0 16px;">
            <span style="color:#64748b; font-size:0.85rem; letter-spacing:0.08em; text-transform:uppercase; font-weight:600;">
                ↓ Presiona el micrófono ↓
            </span>
        </div>
        """, unsafe_allow_html=True)

        rec_c1, rec_c2, rec_c3 = st.columns([5, 1, 5])
        with rec_c2:
            st.markdown('<div style="display:flex; flex-direction:column; align-items:center; justify-content:center; margin-bottom: 24px;">', unsafe_allow_html=True)
            audio_bytes = audio_recorder(
                text="",
                recording_color="#a78bfa",
                neutral_color="#4f46e5",
                icon_size="3x",
                pause_threshold=3.0,
            )
            st.markdown('<p style="color:#e2e8f0; font-size:1rem; font-weight:600; margin-top:8px;">Grabar</p>', unsafe_allow_html=True)
            st.markdown('</div>', unsafe_allow_html=True)



        if audio_bytes:
            st.audio(audio_bytes, format="audio/wav")
            st.markdown("")

            col_l2, col_btn2, col_r2 = st.columns([1, 2, 1])
            with col_btn2:
                rec_btn = st.button("🚀 Transcribir Grabación", use_container_width=True, key="btn_record")

            if rec_btn:
                with st.spinner("Cargando modelo de IA…"):
                    model = load_whisper_model(model_choice)

                with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as tmp:
                    tmp.write(audio_bytes)
                    tmp_path = tmp.name

                try:
                    text, detected = transcribe_audio(tmp_path, model, selected_lang)
                    save_transcription("grabacion_microfono.wav", detected, model_choice, text)
                    st.session_state["last_text"]     = text
                    st.session_state["last_filename"] = "grabacion_microfono.wav"
                    st.session_state["last_lang"]     = detected
                    st.success(f"✅ Transcripción guardada · Idioma detectado: **{detected}**")
                except Exception as e:
                    st.error(f"❌ Error durante la transcripción: {e}")
                finally:
                    if os.path.exists(tmp_path):
                        os.remove(tmp_path)

    except ImportError:
        st.warning(
            "⚠️ Para habilitar la grabación de micrófono, instala el paquete:\n\n"
            "```bash\npip install audio-recorder-streamlit\n```"
        )


# ══════════════════════════════════════════════════════════════════════════════
# TAB 3 — History
# ══════════════════════════════════════════════════════════════════════════════
with tab_history:
    st.markdown("")
    history = load_history()

    if not history:
        st.info("📭 Aún no tienes transcripciones guardadas. ¡Sube un archivo o graba tu voz!")
    else:
        for row in history:
            row_id, filename, language, model_used, text, word_count, created_at = row
            preview = text[:120] + "…" if len(text) > 120 else text

            with st.expander(f"**{filename}** — {created_at}", expanded=False):
                col_meta1, col_meta2, col_meta3 = st.columns(3)
                with col_meta1:
                    st.metric("Palabras", f"{word_count:,}")
                with col_meta2:
                    st.metric("Idioma", language.upper() if language else "—")
                with col_meta3:
                    st.metric("Modelo", model_used)

                edited = st.text_area(
                    "Transcripción",
                    value=text,
                    height=200,
                    key=f"hist_text_{row_id}"
                )

                dl_col1, dl_col2, dl_col3, del_col = st.columns([2, 2, 2, 1])
                with dl_col1:
                    st.download_button(
                        "📄 TXT",
                        data=export_txt(edited),
                        file_name=f"{Path(filename).stem}.txt",
                        mime="text/plain",
                        use_container_width=True,
                        key=f"dl_txt_{row_id}"
                    )
                with dl_col2:
                    st.download_button(
                        "📝 DOCX",
                        data=export_docx(edited, filename),
                        file_name=f"{Path(filename).stem}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        key=f"dl_docx_{row_id}"
                    )
                with dl_col3:
                    st.download_button(
                        "📊 PDF",
                        data=export_pdf(edited, filename),
                        file_name=f"{Path(filename).stem}.pdf",
                        mime="application/pdf",
                        use_container_width=True,
                        key=f"dl_pdf_{row_id}"
                    )
                with del_col:
                    if st.button("🗑️", key=f"del_{row_id}", help="Eliminar transcripción"):
                        delete_transcription(row_id)
                        st.rerun()


# ══════════════════════════════════════════════════════════════════════════════
# Result Panel (shown below tabs after any transcription)
# ══════════════════════════════════════════════════════════════════════════════
if st.session_state.get("last_text"):
    st.markdown("---")
    st.markdown("### 📝 Última Transcripción")

    text_val    = st.session_state["last_text"]
    fname_val   = st.session_state.get("last_filename", "audio")
    lang_val    = st.session_state.get("last_lang", "—")
    words_val   = len(text_val.split())
    chars_val   = len(text_val)

    m1, m2, m3 = st.columns(3)
    m1.metric("🌐 Idioma detectado", lang_val.upper())
    m2.metric("📊 Palabras",         f"{words_val:,}")
    m3.metric("🔡 Caracteres",        f"{chars_val:,}")

    edited_result = st.text_area(
        "Puedes editar la transcripción aquí:",
        value=text_val,
        height=280,
        key="result_editor"
    )

    st.markdown("#### 💾 Exportar")
    ec1, ec2, ec3 = st.columns(3)
    with ec1:
        st.download_button(
            "📄 Descargar TXT",
            data=export_txt(edited_result),
            file_name=f"{Path(fname_val).stem}.txt",
            mime="text/plain",
            use_container_width=True,
            key="exp_txt"
        )
    with ec2:
        st.download_button(
            "📝 Descargar DOCX",
            data=export_docx(edited_result, fname_val),
            file_name=f"{Path(fname_val).stem}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key="exp_docx"
        )
    with ec3:
        st.download_button(
            "📊 Descargar PDF",
            data=export_pdf(edited_result, fname_val),
            file_name=f"{Path(fname_val).stem}.pdf",
            mime="application/pdf",
            use_container_width=True,
            key="exp_pdf"
        )

    # ── Sección Ollama IA ──────────────────────────────────────────────────────
    st.markdown("---")
    st.markdown("""
    <div style="display:flex; align-items:center; gap:10px; margin-bottom:8px;">
        <span style="font-size:1.3rem;">🤖</span>
        <span style="font-size:1.1rem; font-weight:700; color:#f1f5f9;">Mejorar con IA</span>
        <span style="
            background: rgba(124,58,237,0.15);
            border: 1px solid rgba(124,58,237,0.4);
            border-radius: 999px;
            padding: 2px 10px;
            font-size: 0.72rem;
            font-weight: 700;
            color: #a78bfa;
            letter-spacing: 0.08em;
        ">OLLAMA · llama3.2</span>
    </div>
    <p style="color:#64748b; font-size:0.88rem; margin-bottom:12px;">
        Usa IA local para mejorar, resumir o transformar tu transcripción.
    </p>
    """, unsafe_allow_html=True)

    ai_col1, ai_col2 = st.columns([2, 1])
    with ai_col1:
        ai_action = st.selectbox(
            "Acción de IA",
            list(OLLAMA_PROMPTS.keys()),
            label_visibility="collapsed"
        )
    with ai_col2:
        run_ai = st.button("▶ Procesar con IA", use_container_width=True, key="btn_ollama")

    if run_ai:
        with st.spinner(f"🤖 Procesando con Ollama ({OLLAMA_MODEL})…"):
            ai_result = run_ollama(OLLAMA_PROMPTS[ai_action], edited_result)
        st.session_state["ai_result"] = ai_result
        st.session_state["ai_action"] = ai_action

    if st.session_state.get("ai_result"):
        st.markdown(f"**Resultado — {st.session_state.get('ai_action', '')}:**")
        
        ai_edited = st.text_area(
            "Resultado IA (puedes editarlo)",
            value=st.session_state["ai_result"],
            height=220,
            key="ai_editor",
            label_visibility="collapsed"
        )
        
        st.caption("📋 *Para copiar rápidamente, haz clic en el botón de la esquina superior derecha del siguiente recuadro:*")
        st.code(ai_edited, language="markdown")

        ai_d1, ai_d2, ai_d3, ai_d4 = st.columns(4)
        with ai_d1:
            st.download_button("📄 TXT", export_txt(ai_edited),
                               f"{Path(fname_val).stem}_ia.txt", "text/plain",
                               use_container_width=True, key="ai_txt")
        with ai_d2:
            st.download_button("📝 DOCX", export_docx(ai_edited, fname_val),
                               f"{Path(fname_val).stem}_ia.docx",
                               "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                               use_container_width=True, key="ai_docx")
        with ai_d3:
            st.download_button("📊 PDF", export_pdf(ai_edited, fname_val),
                               f"{Path(fname_val).stem}_ia.pdf", "application/pdf",
                               use_container_width=True, key="ai_pdf")
        with ai_d4:
            if st.button("↩ Usar como transcripción", use_container_width=True, key="ai_replace"):
                st.session_state["last_text"] = ai_edited
                st.rerun()